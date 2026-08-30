"""
Извлечение норм СТО-ССР-2026 «Единичные нормы для планирования ресурсов
в ПО Spider Project» — внутренний стандарт подрядчика ООО «ССР» (той же
организации, что ведёт ТМ-35), реально применялся Заказчиком при вводе
отчётов с объекта в Spider Project (практика позже прекращена — эти два
файла единственное, что осталось). Координатор подтвердил: Spider
Project, по всей видимости, подключён к ГЭСН через API, поэтому нормы
местами близки к ГЭСН — но раздел работ здесь прямо соответствует
ТМ-35 (Устройство трубопроводов, Устройство камеры, СОДК и т.п.), в
отличие от общего 16-сборникового отбора ГЭСН, сделанного вручную.

Источник — только "версия 4 Приложение 1 ... (2).docx": подтверждено
построчной сверкой, что это строгое расширение более раннего файла
(та же 370 операций + доп. строка "Трудоёмкость людских ресурсов").

Структура документа (python-docx, элементы body в порядке следования):
  параграф "(КОД) Название операции" -> раздел контекста (последний
  встреченный заголовок раздела вида "Название (АББР)") -> далее
  несколько служебных параграфов ("Указание по применению норм",
  "Состав работы", список пунктов, "Тип ДПГ - ...") -> таблица:
    шапка 1: "№п/п | Наименование работ | Состав команды (x4) | Производительность одного ресурса в (ЕД/час)"
    шапка 2: "№п/п | Наименование работ | Ресурс | Ед.изм. | Кол-во | Загрузка(%) | Производительность одного ресурса в (ЕД/час)"
    строки 1.1..1.N — по одному ресурсу команды (produktivnost непустая только у ключевого ресурса)
    строка "Производительность на команду/1 бригаду" — итоговая произв-ть бригады
    строка "Трудоёмкость людских ресурсов" — чел-час/ед. (если есть)
"""
import json
import re

import docx

CODE_RE = re.compile(r"^\((\S+)\)\s*(.+)$")
SECTION_RE = re.compile(r"^([А-ЯЁа-яё][А-Яа-яё \-,]{3,60})\(([А-ЯЁ]{2,6})\)$")
UNIT_RE = re.compile(r"Производительность одного ресурса в \(([^/]+)/час\)")


def parse(path):
    d = docx.Document(path)
    body = d.element.body

    # Быстрый доступ: параграф/таблица по объекту XML -> python-docx обёртка
    para_by_elem = {p._p: p for p in d.paragraphs}
    table_by_elem = {t._tbl: t for t in d.tables}

    operations = []
    current_section = None
    current_code = None
    current_name = None
    pending_note_lines = []

    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            p = para_by_elem.get(child)
            if p is None:
                continue
            text = p.text.strip()
            if not text:
                continue
            sm = SECTION_RE.match(text)
            cm = CODE_RE.match(text)
            if cm:
                current_code, current_name = cm.group(1), cm.group(2).strip()
                pending_note_lines = []
            elif sm and not cm:
                current_section = f"{sm.group(1).strip()} ({sm.group(2)})"
            elif current_code and text not in ("Указание по применению норм", "Состав работы", "Состав работ") \
                    and not text.startswith("Тип ДПГ"):
                pending_note_lines.append(text)
        elif child.tag.endswith("}tbl"):
            t = table_by_elem.get(child)
            if t is None or current_code is None:
                continue
            op = parse_table(t, current_section, current_code, current_name)
            if op:
                op["notes"] = " ".join(pending_note_lines[-6:])  # состав работы, последние строки перед таблицей
                operations.append(op)
            current_code = None  # таблица «использована» — до следующего кода не переоткрывать

    return operations


def _num(s):
    s = (s or "").strip().replace(" ", "").replace(",", ".")
    if not s:
        return None
    try:
        return float(s)
    except ValueError:
        return None


def parse_table(table, section, code, name):
    rows = [[c.text.strip() for c in r.cells] for r in table.rows]
    if len(rows) < 2:
        return None

    unit = None
    m = UNIT_RE.search(rows[0][-1]) if rows[0] else None
    if m:
        unit = m.group(1).strip()

    crew = []
    team_productivity = None
    labor_hours = None

    for row in rows[2:]:
        if len(row) < 3:
            continue
        label = row[1].strip()
        if label.startswith("Производительность на"):
            team_productivity = _num(row[-1])
            continue
        if label.startswith("Трудоёмкость"):
            labor_hours = _num(row[-1])
            continue
        # обычная строка ресурса: №пп | наименование | ресурс | ед | кол-во | загрузка% | произв-ть
        if len(row) >= 7:
            resource, res_unit, qty, loading, prod = row[2], row[3], row[4], row[5], row[6]
        elif len(row) == 6:
            resource, res_unit, qty, loading, prod = row[2], row[3], row[4], row[5], ""
        else:
            continue
        if not resource:
            continue
        crew.append({
            "resource": resource.strip(),
            "unit": res_unit.strip(),
            "qty": _num(qty),
            "loading_pct": _num(loading),
            "resource_productivity": _num(prod),
        })

    if not crew and team_productivity is None:
        return None

    return {
        "section": section,
        "code": code,
        "name": name,
        "unit": unit,
        "crew": crew,
        "team_productivity_per_hour": team_productivity,
        "labor_hours_per_unit": labor_hours,
    }


def main():
    path = "/home/oleg/Downloads/версия 4 Приложение 1 Единичные нормы для планирования ресурсов в ПО Spider Project (2).docx"
    ops = parse(path)

    out_path = "/home/oleg/Documents/TM-35/import/enir_work/ssr_spider_norms.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(ops, f, ensure_ascii=False, indent=1)

    n_total = len(ops)
    n_with_labor = sum(1 for o in ops if o["labor_hours_per_unit"] is not None)
    n_with_prod = sum(1 for o in ops if o["team_productivity_per_hour"] is not None)
    by_section = {}
    for o in ops:
        by_section.setdefault(o["section"], [0, 0])
        by_section[o["section"]][0] += 1
        if o["labor_hours_per_unit"] is not None:
            by_section[o["section"]][1] += 1

    print(f"Всего операций: {n_total}")
    print(f"С производительностью бригады: {n_with_prod}")
    print(f"С трудоёмкостью чел-час/ед: {n_with_labor}")
    print(f"\n{'Раздел':45s} {'опер.':>6s} {'с трудоёмк.':>12s}")
    for sec, (total, ok) in sorted(by_section.items(), key=lambda x: str(x[0])):
        print(f"{str(sec):45s} {total:>6d} {ok:>12d}")


if __name__ == "__main__":
    main()
